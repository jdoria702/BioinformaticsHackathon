from __future__ import annotations

import os
import logging

from typing import List, Tuple, Dict, Optional

from PIL import Image
import pytesseract
from pypdf import PdfReader
import docx
import base64
import requests

logger = logging.getLogger(__name__)

SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}

# Extract text from PDF uploads:
def extract_pdf_blocks(path: str) -> List[str]:
    logger.info("Extracting text from PDF at %s...", path)

    # Initialize PDF reader:
    reader = PdfReader(path)
    parts: List[str] = []

    # Split content page by page:
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()

        if text:
            parts.append(text)
        else:
            logger.debug("No extractable text found on PDF page %d", page_number)
    
    logger.info("Extracted %d text blocks from PDF", len(parts))
    return parts

# Extract text + tables from docx:
def extract_docx_blocks(path: str) -> List[str]:
    logger.info("Extracting text from .docx at %s...", path)

    # Split content by paragraph:
    document = docx.Document(path)
    parts: List[str] = []

    # Extract paragraph text:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Extract table text:
    # Table-by-table:
    for table in document.tables:
        # Row-by-row:
        for row in table.rows:
            row_text: List[str] = []
            # Cell-by-cell:
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)

            # If row is empty -> delete:
            has_content = False

            for cell in row_text:
                if cell:
                    has_content = True
                    break
        
            if has_content:
                parts.append(" | ".join(row_text))

    logger.info("Extracted %d text blocks from .docx", len(parts))
    return parts

# Extract text from images using OCR:
def extract_image_ocr_text(path: str) -> str:
    logger.info("Running OCR on image at %s...", path)
    try:
        image = Image.open(path)
        return pytesseract.image_to_string(image).strip()
    except Exception:
        logger.exception("OCR failed for image %s", path)
        return ""

# Describe an image and store the description using a multimodal model:
def describe_image_with_model(path: str) -> str:
    """Uses Ollama + Moondream to generate an image description."""

    logger.info("Generating image description for %s...", path)

    try:
        # Read image and encode as base64:
        with open(path, "rb") as f:
            image_bytes = f.read()
            image_base64 = base64.b64encode(image_bytes).decode("utf-8")

        # Ollama API request:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "moondream",
                "prompt": (
                    "Describe this image clearly and concisely for retrieval. "
                    "Focus on key objects, labels, and scientific meaning if present."
                ),
                "images": [image_base64],
                "stream": False,
            },
            timeout=60,
        )

        response.raise_for_status()
        result = response.json()

        description = result.get("response", "").strip()

        logger.debug("Image description: %s", description)

        return description

    except requests.exceptions.ConnectionError:
        logger.warning("Ollama not reachable at http://localhost:11434 (skipping image captioning)")
        return ""
    except Exception:
        logger.exception("Failed to generate image description for %s", path)
        return ""

# Helper function to extract image info in blocks:
def extract_image_blocks(
    path: str,
    use_ocr: bool = True,
    use_caption: bool = True,
) -> List[str]:
    logger.info(
        "Extracting image blocks from %s (ocr=%s, caption=%s)",
        path,
        use_ocr,
        use_caption,
    )

    parts: List[str] = []

    # OCR extraction:
    if use_ocr:
        ocr_text = extract_image_ocr_text(path)
        if ocr_text:
            parts.append(f"OCR: {ocr_text}")
    
    # Vision captioning extraction:
    if use_caption:
        caption = describe_image_with_model(path)
        if caption:
            parts.append(f"IMAGE_DESCRIPTION: {caption}")

    if not parts:
        logger.warning("No text extracted from image %s (ocr=%s, caption=%s)", path, use_ocr, use_caption)

    logger.info("Extracted %d text blocks from image", len(parts))
    return parts

# Helper to build block metadata:
def build_block_metadata(
    path: str,
    kind: str,
    block_count: int,
    ingestion_method: Optional[str] = None
) -> List[Dict]:
    source = os.path.basename(path)
    metadatas: List[Dict] = []

    for index in range(block_count):
        metadata: Dict = {
            "source": source,
            "file_type": kind,
            "block_index": index,
        }

        if ingestion_method is not None:
            metadata["ingestion_method"] = ingestion_method
        
        metadatas.append(metadata)
    
    return metadatas

# Helper function to extract blocks:
def extract_blocks(path: str) -> Tuple[List[str], str]:
    """
    Returns (blocks, kind).
    - Blocks is always a list of text blocks.
    - Kind is either a .pdf, .docx, or image file type.
    """
    extension = os.path.splitext(path.lower())[1]

    # .pdf:
    if extension == ".pdf":
        return extract_pdf_blocks(path), "pdf"

    # .docx:
    if extension == ".docx":
        return extract_docx_blocks(path), "docx"
    
    # Images:
    if extension in SUPPORTED_IMAGE_EXTENSIONS:
        return extract_image_blocks(path), "image"
    
    raise ValueError(f"Unsupported file type: {extension}")

# Chunk text:
def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> List[str]:
    text = (text or "").strip()

    if not text:
        return []
    
    if max_chars <= 0:
        raise ValueError("max_chars must be greater than 0.")
    
    if overlap < 0:
        raise ValueError("overlap cannot be negative.")
    
    if overlap >= max_chars:
        raise ValueError("overlap must be smaller than max_chars.")    
    
    chunks: List[str] = []
    start = 0

    while start < len(text):
        end = min(len(text), start + max_chars)
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end == len(text):
            break

        start = end - overlap

    return chunks

# Chunk blocks:
def chunk_blocks(blocks: List[str], max_chars: int = 1200, overlap: int = 150) -> List[str]:
    """
    Chunk a list of blocks into smaller chunks.
    Short blocks stay as-is; long blocks get split with overlap.
    """
    chunks: List[str] = []

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        
        if len(block) <= max_chars:
            chunks.append(block)
        else:
            chunks.extend(chunk_text(block, max_chars=max_chars, overlap=overlap))

    return chunks
